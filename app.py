# ==========================================
# app.py
# ==========================================
import streamlit as st
import streamlit.components.v1 as components
import jarvis_config as cfg
import jarvis_engine as engine
from docx import Document
from gtts import gTTS
import os
import io
import re
import base64

# ==========================================
# 背景處理函數
# ==========================================
def save_memory_to_word(tags, memories, goals):
    try:
        doc = Document()
        doc.add_heading('Jarvis Memory Log', 0)
        doc.add_heading('24 維度標籤', level=1)
        doc.add_paragraph(tags if tags else "無")
        doc.add_heading('專屬記憶', level=1)
        doc.add_paragraph(memories if memories else "無")
        doc.add_heading('目標庫存', level=1)
        doc.add_paragraph(goals if goals else "無")
        
        file_path = "jarvis_memory_log.docx"
        doc.save(file_path)
    except Exception as e:
        pass

def generate_audio_bytes(text):
    """將文字轉換為語音並返回 Byte 數據"""
    clean_text = re.sub(r'[*_#`~]', '', text)
    if not clean_text.strip(): return None
    try:
        tts = gTTS(text=clean_text, lang='zh-tw')
        fp = io.BytesIO()
        tts.write_to_fp(fp)
        fp.seek(0)
        return fp.read()
    except:
        return None

def render_audio_player(audio_bytes, speed=1.8, autoplay=False):
    """利用前端 HTML/JS 強制改變語音播放速度 (預設 1.8 倍)"""
    if not audio_bytes: return
    b64 = base64.b64encode(audio_bytes).decode()
    auto_attr = "autoplay" if autoplay else ""
    html_code = f"""
        <audio id="jarvis_audio_{str(hash(audio_bytes))[-5:]}" controls {auto_attr} style="width: 100%; height: 40px;">
            <source src="data:audio/mp3;base64,{b64}" type="audio/mp3">
        </audio>
        <script>
            var audio = document.getElementById("jarvis_audio_{str(hash(audio_bytes))[-5:]}");
            audio.playbackRate = {speed};
        </script>
    """
    components.html(html_code, height=50)

# ==========================================
# 1. 頁面與狀態初始化
# ==========================================
st.set_page_config(page_title="Jarvis Command Center", layout="wide", initial_sidebar_state="expanded")

if "messages" not in st.session_state:
    st.session_state.messages = []
if "available_models" not in st.session_state:
    st.session_state.available_models = []

if "state_tags" not in st.session_state:
    st.session_state.state_tags = cfg.DEFAULT_24_TAGS
if "state_memories" not in st.session_state:
    st.session_state.state_memories = ""
if "state_goals" not in st.session_state:
    st.session_state.state_goals = ""
if "state_scores" not in st.session_state:
    st.session_state.state_scores = ""

# ==========================================
# 2. 側邊欄：API 與模型鎖定
# ==========================================
with st.sidebar:
    st.title("⚙️ Jarvis 系統控制")
    api_key = st.text_input("🔑 API 金鑰", value=cfg.DEFAULT_API_KEY, type="password")
    
    if api_key:
        if st.button("🔄 重新整理可用模型清單") or not st.session_state.available_models:
            with st.spinner("正在向 Google 請求可用模型..."):
                try:
                    st.session_state.available_models = engine.fetch_available_models(api_key)
                except Exception as e:
                    st.error(f"無法獲取清單: {e}")

        if st.session_state.available_models:
            default_idx = 0
            # 強制尋找 3.1-pro 或 pro-preview
            for i, m in enumerate(st.session_state.available_models):
                if "3.1-pro" in m.lower() or "pro-preview" in m.lower():
                    default_idx = i
                    break
            
            selected_model = st.selectbox("🤖 選擇運算核心 (Model)", st.session_state.available_models, index=default_idx)
            st.info(f"當前模型：{selected_model}")
            
    st.markdown("---")
    st.subheader("🗣️ 語音設定")
    playback_speed = st.slider("朗讀倍速", min_value=1.0, max_value=2.5, value=1.8, step=0.1)

    st.markdown("---")
    st.markdown("### 📦 模組說明速查")
    category = st.selectbox("選擇模組分類", list(cfg.MODULES_FOR_UI.keys()))
    for mod_name, mod_desc in cfg.MODULES_FOR_UI[category].items():
        with st.expander(f"🔹 {mod_name}"):
            st.caption(mod_desc)
            
    st.markdown("---")
    if os.path.exists("jarvis_memory_log.docx"):
        with open("jarvis_memory_log.docx", "rb") as file:
            st.download_button(
                label="💾 下載最新 Memory Word 檔",
                data=file,
                file_name="jarvis_memory_log.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# ==========================================
# 3. 雙欄式主畫面：左側對話區 / 右側即時分析板
# ==========================================
col_chat, col_dash = st.columns([7, 3], gap="large")

with col_chat:
    st.title("Jarvis 終端控制台")
    
    for msg in st.session_state.messages:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])
            if msg.get("audio_bytes"):
                render_audio_player(msg["audio_bytes"], speed=playback_speed, autoplay=False)

    if user_input := st.chat_input("輸入指令，先生..."):
        if not api_key:
            st.error("先生，請先配置 API Key。")
            st.stop()
        
        st.session_state.messages.append({"role": "user", "content": user_input})
        with st.chat_message("user"):
            st.markdown(user_input)

        with st.chat_message("assistant"):
            with st.spinner(f'Jarvis ({selected_model}) 戰略推演中...'):
                try:
                    history_for_api = []
                    for m in st.session_state.messages[:-1]:
                        if m["role"] == "user":
                            history_for_api.append({"role": "user", "parts": [m["content"]]})
                        else:
                            full_memory = m.get("raw_text", m["content"])
                            history_for_api.append({"role": "model", "parts": [full_memory]})
                            
                    forced_input = cfg.get_forced_template(
                        user_input=user_input, 
                        tags=st.session_state.state_tags, 
                        memories=st.session_state.state_memories, 
                        goals=st.session_state.state_goals, 
                        scores=st.session_state.state_scores
                    )
                    
                    result = engine.process_jarvis_turn(
                        api_key=api_key,
                        selected_model=selected_model,
                        system_prompt=cfg.SYSTEM_PROMPT,
                        history_for_api=history_for_api,
                        forced_template_text=forced_input
                    )
                    
                    d = result["parsed_dash"]
                    
                    # 覆寫 24 項標籤
                    if d.get("tags") and "未解析" not in d.get("tags"):
                        st.session_state.state_tags = d["tags"]
                        
                    # 增量累積專屬記憶
                    new_mem = d.get("new_memory", "")
                    if new_mem and "未解析" not in new_mem and "無" not in new_mem:
                        clean_mem = new_mem.replace("[新增]", "").strip()
                        if clean_mem:
                            st.session_state.state_memories += f"\n- {clean_mem}"
                            
                    # 更新目標
                    if d.get("new_goal") and "未解析" not in d.get("new_goal"):
                        st.session_state.state_goals = d["new_goal"]
                        
                    st.session_state.state_scores = f"友善度:{d.get('friendly', 'N/A')}, 信任度:{d.get('trust', 'N/A')}, SAI:{d.get('sai', 'N/A')}, 準確度:{d.get('accuracy', 'N/A')}"
                    
                    save_memory_to_word(st.session_state.state_tags, st.session_state.state_memories, st.session_state.state_goals)
                    
                    st.markdown(result["output"])
                    
                    audio_b = generate_audio_bytes(result["output"])
                    if audio_b:
                        render_audio_player(audio_b, speed=playback_speed, autoplay=True)
                    
                    st.session_state.messages.append({
                        "role": "assistant",
                        "raw_text": result["raw_full_text"],     
                        "content": result["output"],
                        "parsed_dash": result["parsed_dash"],
                        "audio_bytes": audio_b
                    })
                    st.rerun() 

                except Exception as e:
                    st.error(f"運算中斷：{str(e)}")

# ==========================================
# 4. 右側欄：即時戰略分析板 (Dashboard)
# ==========================================
with col_dash:
    st.subheader("📊 實時動態分析板")
    st.markdown("*(擷取自最新一輪 AI 運算結果)*")
    st.divider()
    
    latest_msg = None
    for msg in reversed(st.session_state.messages):
        if msg["role"] == "assistant":
            latest_msg = msg
            break
            
    if latest_msg and latest_msg.get("parsed_dash"):
        d = latest_msg["parsed_dash"]
        
        st.markdown("**1. 啟用模組 (激活)**")
        st.info(d.get("modules", "無"))
        
        st.markdown("**2. 標籤與特質庫存**")
        with st.expander("24 維度標籤", expanded=False):
            st.code(st.session_state.state_tags)
        with st.expander("專屬記憶 (累積)", expanded=False):
            st.write(st.session_state.state_memories if st.session_state.state_memories else "無")
        
        st.markdown("**3. 意圖判讀及應對策略 A**")
        st.write(d.get("intent", "無"))
        
        st.markdown("**4. 儀表板變數**")
        with st.expander(f"友善度: {d.get('friendly', 'N/A').split('(')[0]}"):
            st.write(d.get("friendly", "無資料"))
        with st.expander(f"信任度: {d.get('trust', 'N/A').split('(')[0]}"):
            st.write(d.get("trust", "無資料"))
        with st.expander(f"SAI 優勢: {d.get('sai', 'N/A').split('(')[0]}"):
            st.write(d.get("sai", "無資料"))
        with st.expander(f"準確度: {d.get('accuracy', 'N/A').split('(')[0]}"):
            st.write(d.get("accuracy", "無資料"))
            
        st.markdown("**5. SAI 修正策略**")
        with st.expander(d.get("sai_strategy", "無").split('判讀')[0]):
            st.write("**判讀理由：**", d.get("sai_reason", "無資料"))
            
        st.markdown("**6. 偽裝與準確度連動矩陣**")
        with st.expander(d.get("matrix", "無")):
            st.write("**判讀理由：**", d.get("matrix_reason", "無資料"))
            
        st.markdown("**7. 產生策略 B**")
        st.write(d.get("strategy_b", "無"))
        
        st.markdown("**8. 融合決策**")
        st.success(d.get("fusion", "無"))
        
        st.markdown("**9. 新目標 (D) / 目標庫存**")
        st.write(d.get("new_goal", "無"))
        
        st.markdown("**10. 決定次輪策略 (D)**")
        st.warning(d.get("next_strategy", "無"))
        
        st.divider()
        st.caption("⚙️ 開發者底層監控")
        with st.expander("🔍 展開底層原始運算 Log (Raw Data)", expanded=False):
            st.code(latest_msg.get("raw_text", "無資料"), language="markdown")
        
    else:
        st.caption("等待首輪對話產生運算結果...")
